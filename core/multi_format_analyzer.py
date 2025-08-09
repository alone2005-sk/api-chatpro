"""
Multi-Format File Analyzer
Handles PDF, Word, Excel, audio, video, images, and other file formats
"""

import os
import asyncio
import tempfile
import logging
from pathlib import Path
from typing import Dict, Any, List, Optional, Union
import mimetypes

# File processing imports
import PyPDF2
import docx
import pandas as pd
import whisper
import cv2
from PIL import Image
import json
import xml.etree.ElementTree as ET
import yaml

logger = logging.getLogger(__name__)

class MultiFormatAnalyzer:
    def __init__(self):
        self.temp_dir = Path(tempfile.mkdtemp(prefix="analyzer_"))
        self.whisper_model = None
        
        # Supported file types
        self.supported_formats = {
            'pdf': ['.pdf'],
            'document': ['.doc', '.docx', '.txt', '.rtf'],
            'spreadsheet': ['.xlsx', '.xls', '.csv'],
            'audio': ['.mp3', '.wav', '.m4a', '.flac', '.ogg'],
            'video': ['.mp4', '.avi', '.mov', '.mkv', '.webm'],
            'image': ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff'],
            'code': ['.py', '.js', '.html', '.css', '.java', '.cpp', '.go', '.rs'],
            'data': ['.json', '.xml', '.yaml', '.yml'],
            'archive': ['.zip', '.tar', '.gz', '.rar']
        }
    
    async def initialize(self):
        """Initialize the analyzer with required models"""
        try:
            # Load Whisper model for audio processing
            self.whisper_model = whisper.load_model("base")
            logger.info("Multi-format analyzer initialized with Whisper model")
        except Exception as e:
            logger.warning(f"Could not load Whisper model: {e}")
    
    async def analyze_file(
        self, 
        file, 
        extract_audio: bool = False, 
        extract_images: bool = False
    ) -> Dict[str, Any]:
        """Analyze a single file and extract relevant information"""
        try:
            # Save uploaded file temporarily
            temp_path = self.temp_dir / file.filename
            with open(temp_path, 'wb') as f:
                content = await file.read()
                f.write(content)
            
            # Detect file type
            file_type = self._detect_file_type(temp_path)
            
            # Process based on file type
            result = {
                "filename": file.filename,
                "file_type": file_type,
                "size": len(content),
                "mime_type": mimetypes.guess_type(file.filename)[0]
            }
            
            if file_type == 'pdf':
                result.update(await self._analyze_pdf(temp_path))
            elif file_type == 'document':
                result.update(await self._analyze_document(temp_path))
            elif file_type == 'spreadsheet':
                result.update(await self._analyze_spreadsheet(temp_path))
            elif file_type == 'audio':
                result.update(await self._analyze_audio(temp_path))
            elif file_type == 'video':
                result.update(await self._analyze_video(temp_path, extract_audio, extract_images))
            elif file_type == 'image':
                result.update(await self._analyze_image(temp_path))
            elif file_type == 'code':
                result.update(await self._analyze_code(temp_path))
            elif file_type == 'data':
                result.update(await self._analyze_data(temp_path))
            else:
                result.update(await self._analyze_generic(temp_path))
            
            # Cleanup
            temp_path.unlink()
            
            return result
            
        except Exception as e:
            logger.error(f"Error analyzing file {file.filename}: {str(e)}")
            return {
                "filename": file.filename,
                "error": str(e),
                "file_type": "unknown"
            }
    
    def _detect_file_type(self, file_path: Path) -> str:
        """Detect file type based on extension"""
        extension = file_path.suffix.lower()
        
        for file_type, extensions in self.supported_formats.items():
            if extension in extensions:
                return file_type
        
        return "unknown"
    
    async def _analyze_pdf(self, file_path: Path) -> Dict[str, Any]:
        """Extract text and metadata from PDF files"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Extract metadata
                metadata = pdf_reader.metadata or {}
                
                # Extract text from all pages
                text_content = []
                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    text_content.append({
                        "page": page_num + 1,
                        "text": page_text,
                        "word_count": len(page_text.split())
                    })
                
                # Combine all text
                full_text = " ".join([page["text"] for page in text_content])
                
                return {
                    "pages": len(pdf_reader.pages),
                    "text_content": text_content,
                    "full_text": full_text,
                    "word_count": len(full_text.split()),
                    "metadata": {
                        "title": metadata.get("/Title", ""),
                        "author": metadata.get("/Author", ""),
                        "subject": metadata.get("/Subject", ""),
                        "creator": metadata.get("/Creator", "")
                    }
                }
                
        except Exception as e:
            return {"error": f"PDF analysis failed: {str(e)}"}
    
    async def _analyze_document(self, file_path: Path) -> Dict[str, Any]:
        """Analyze Word documents and text files"""
        try:
            if file_path.suffix.lower() in ['.doc', '.docx']:
                # Word document
                doc = docx.Document(file_path)
                
                paragraphs = []
                for para in doc.paragraphs:
                    if para.text.strip():
                        paragraphs.append(para.text)
                
                full_text = "\n".join(paragraphs)
                
                # Extract tables if any
                tables = []
                for table in doc.tables:
                    table_data = []
                    for row in table.rows:
                        row_data = [cell.text for cell in row.cells]
                        table_data.append(row_data)
                    tables.append(table_data)
                
                return {
                    "paragraphs": paragraphs,
                    "full_text": full_text,
                    "word_count": len(full_text.split()),
                    "paragraph_count": len(paragraphs),
                    "tables": tables,
                    "table_count": len(tables)
                }
            else:
                # Plain text file
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                
                return {
                    "content": content,
                    "word_count": len(content.split()),
                    "line_count": len(content.splitlines()),
                    "character_count": len(content)
                }
                
        except Exception as e:
            return {"error": f"Document analysis failed: {str(e)}"}
    
    async def _analyze_spreadsheet(self, file_path: Path) -> Dict[str, Any]:
        """Analyze Excel and CSV files"""
        try:
            if file_path.suffix.lower() == '.csv':
                df = pd.read_csv(file_path)
            else:
                df = pd.read_excel(file_path, sheet_name=None)
                
                # If multiple sheets, analyze each
                if isinstance(df, dict):
                    sheets_info = {}
                    for sheet_name, sheet_df in df.items():
                        sheets_info[sheet_name] = {
                            "rows": len(sheet_df),
                            "columns": len(sheet_df.columns),
                            "column_names": list(sheet_df.columns),
                            "sample_data": sheet_df.head().to_dict('records')
                        }
                    
                    return {
                        "sheet_count": len(df),
                        "sheets": sheets_info,
                        "total_rows": sum(len(sheet_df) for sheet_df in df.values())
                    }
                else:
                    df = df
            
            # Single sheet analysis
            return {
                "rows": len(df),
                "columns": len(df.columns),
                "column_names": list(df.columns),
                "data_types": df.dtypes.to_dict(),
                "sample_data": df.head().to_dict('records'),
                "summary_stats": df.describe().to_dict() if df.select_dtypes(include='number').columns.any() else None
            }
            
        except Exception as e:
            return {"error": f"Spreadsheet analysis failed: {str(e)}"}
    
    async def _analyze_audio(self, file_path: Path) -> Dict[str, Any]:
        """Transcribe audio files using Whisper"""
        try:
            if not self.whisper_model:
                return {"error": "Whisper model not available"}
            
            # Transcribe audio
            result = self.whisper_model.transcribe(str(file_path))
            
            # Extract segments with timestamps
            segments = []
            for segment in result.get("segments", []):
                segments.append({
                    "start": segment["start"],
                    "end": segment["end"],
                    "text": segment["text"]
                })
            
            return {
                "transcription": result["text"],
                "language": result.get("language", "unknown"),
                "segments": segments,
                "duration": segments[-1]["end"] if segments else 0,
                "word_count": len(result["text"].split())
            }
            
        except Exception as e:
            return {"error": f"Audio analysis failed: {str(e)}"}
    
    async def _analyze_video(
        self, 
        file_path: Path, 
        extract_audio: bool = False, 
        extract_images: bool = False
    ) -> Dict[str, Any]:
        """Analyze video files - extract audio, frames, metadata"""
        try:
            cap = cv2.VideoCapture(str(file_path))
            
            # Get video properties
            fps = cap.get(cv2.CAP_PROP_FPS)
            frame_count = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))
            duration = frame_count / fps if fps > 0 else 0
            width = int(cap.get(cv2.CAP_PROP_FRAME_WIDTH))
            height = int(cap.get(cv2.CAP_PROP_FRAME_HEIGHT))
            
            result = {
                "duration": duration,
                "fps": fps,
                "frame_count": frame_count,
                "resolution": f"{width}x{height}",
                "width": width,
                "height": height
            }
            
            # Extract frames if requested
            if extract_images:
                frames = []
                frame_interval = max(1, frame_count // 10)  # Extract 10 frames max
                
                for i in range(0, frame_count, frame_interval):
                    cap.set(cv2.CAP_PROP_POS_FRAMES, i)
                    ret, frame = cap.read()
                    if ret:
                        # Save frame
                        frame_path = self.temp_dir / f"frame_{i}.jpg"
                        cv2.imwrite(str(frame_path), frame)
                        frames.append({
                            "frame_number": i,
                            "timestamp": i / fps,
                            "path": str(frame_path)
                        })
                
                result["extracted_frames"] = frames
            
            cap.release()
            
            # Extract audio if requested
            if extract_audio and self.whisper_model:
                try:
                    # Use ffmpeg to extract audio (requires ffmpeg installed)
                    audio_path = self.temp_dir / f"{file_path.stem}.wav"
                    cmd = f"ffmpeg -i {file_path} -vn -acodec pcm_s16le -ar 16000 {audio_path}"
                    
                    process = await asyncio.create_subprocess_shell(
                        cmd,
                        stdout=asyncio.subprocess.PIPE,
                        stderr=asyncio.subprocess.PIPE
                    )
                    await process.communicate()
                    
                    if audio_path.exists():
                        # Transcribe extracted audio
                        transcription_result = self.whisper_model.transcribe(str(audio_path))
                        result["audio_transcription"] = transcription_result["text"]
                        result["audio_language"] = transcription_result.get("language", "unknown")
                        
                        # Cleanup
                        audio_path.unlink()
                        
                except Exception as e:
                    result["audio_extraction_error"] = str(e)
            
            return result
            
        except Exception as e:
            return {"error": f"Video analysis failed: {str(e)}"}
    
    async def _analyze_image(self, file_path: Path) -> Dict[str, Any]:
        """Analyze image files - extract metadata and basic properties"""
        try:
            with Image.open(file_path) as img:
                # Basic properties
                result = {
                    "format": img.format,
                    "mode": img.mode,
                    "size": img.size,
                    "width": img.width,
                    "height": img.height
                }
                
                # EXIF data if available
                if hasattr(img, '_getexif') and img._getexif():
                    exif_data = img._getexif()
                    result["exif"] = {str(k): str(v) for k, v in exif_data.items()}
                
                # Color analysis
                if img.mode == 'RGB':
                    colors = img.getcolors(maxcolors=256*256*256)
                    if colors:
                        dominant_colors = sorted(colors, reverse=True)[:5]
                        result["dominant_colors"] = [
                            {"color": color, "count": count} 
                            for count, color in dominant_colors
                        ]
                
                return result
                
        except Exception as e:
            return {"error": f"Image analysis failed: {str(e)}"}
    
    async def _analyze_code(self, file_path: Path) -> Dict[str, Any]:
        """Analyze code files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            lines = content.splitlines()
            
            # Basic metrics
            result = {
                "content": content,
                "line_count": len(lines),
                "character_count": len(content),
                "blank_lines": sum(1 for line in lines if not line.strip()),
                "comment_lines": 0,
                "code_lines": 0
            }
            
            # Language-specific analysis
            extension = file_path.suffix.lower()
            comment_patterns = {
                '.py': ['#'],
                '.js': ['//', '/*', '*/'],
                '.java': ['//', '/*', '*/'],
                '.cpp': ['//', '/*', '*/'],
                '.html': ['<!--', '-->'],
                '.css': ['/*', '*/']
            }
            
            if extension in comment_patterns:
                patterns = comment_patterns[extension]
                for line in lines:
                    stripped = line.strip()
                    if any(stripped.startswith(pattern) for pattern in patterns):
                        result["comment_lines"] += 1
                    elif stripped:
                        result["code_lines"] += 1
            
            # Function/class detection (basic)
            if extension == '.py':
                functions = len([line for line in lines if line.strip().startswith('def ')])
                classes = len([line for line in lines if line.strip().startswith('class ')])
                result.update({"functions": functions, "classes": classes})
            
            return result
            
        except Exception as e:
            return {"error": f"Code analysis failed: {str(e)}"}
    
    async def _analyze_data(self, file_path: Path) -> Dict[str, Any]:
        """Analyze structured data files (JSON, XML, YAML)"""
        try:
            extension = file_path.suffix.lower()
            
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            result = {
                "content": content,
                "size": len(content)
            }
            
            if extension == '.json':
                data = json.loads(content)
                result.update({
                    "type": "json",
                    "structure": self._analyze_json_structure(data),
                    "keys": list(data.keys()) if isinstance(data, dict) else None,
                    "length": len(data) if isinstance(data, (list, dict)) else None
                })
            
            elif extension == '.xml':
                root = ET.fromstring(content)
                result.update({
                    "type": "xml",
                    "root_tag": root.tag,
                    "children_count": len(root),
                    "attributes": root.attrib
                })
            
            elif extension in ['.yaml', '.yml']:
                data = yaml.safe_load(content)
                result.update({
                    "type": "yaml",
                    "structure": self._analyze_json_structure(data),
                    "keys": list(data.keys()) if isinstance(data, dict) else None
                })
            
            return result
            
        except Exception as e:
            return {"error": f"Data analysis failed: {str(e)}"}
    
    async def _analyze_generic(self, file_path: Path) -> Dict[str, Any]:
        """Generic file analysis for unknown types"""
        try:
            with open(file_path, 'rb') as f:
                content = f.read()
            
            # Try to decode as text
            try:
                text_content = content.decode('utf-8')
                return {
                    "type": "text",
                    "content": text_content,
                    "character_count": len(text_content),
                    "line_count": len(text_content.splitlines())
                }
            except UnicodeDecodeError:
                return {
                    "type": "binary",
                    "size": len(content),
                    "binary": True
                }
                
        except Exception as e:
            return {"error": f"Generic analysis failed: {str(e)}"}
    
    def _analyze_json_structure(self, data, max_depth=3, current_depth=0):
        """Analyze JSON structure recursively"""
        if current_depth >= max_depth:
            return "..."
        
        if isinstance(data, dict):
            return {
                key: self._analyze_json_structure(value, max_depth, current_depth + 1)
                for key, value in list(data.items())[:10]  # Limit to first 10 keys
            }
        elif isinstance(data, list):
            if data:
                return [self._analyze_json_structure(data[0], max_depth, current_depth + 1)]
            return []
        else:
            return type(data).__name__
    
    async def get_status(self) -> Dict[str, Any]:
        """Get analyzer status"""
        return {
            "supported_formats": self.supported_formats,
            "whisper_available": self.whisper_model is not None,
            "temp_directory": str(self.temp_dir),
            "status": "ready"
        }
