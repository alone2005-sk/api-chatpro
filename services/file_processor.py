"""
Advanced file processing service for DAMN BOT AI System
Supports multiple file formats with intelligent content extraction
"""

import asyncio
import mimetypes
import os
import tempfile
from pathlib import Path
from typing import Dict, Any, List, Optional
import aiofiles
import magic
from PIL import Image
import pytesseract
import cv2
import numpy as np
from pydub import AudioSegment
import speech_recognition as sr
import pandas as pd
import docx
import PyPDF2
import json
import xml.etree.ElementTree as ET
from bs4 import BeautifulSoup
import zipfile
import tarfile

from core.logger import get_logger
from models.requests import FileProcessResult

logger = get_logger(__name__)

class FileProcessor:
    """Advanced file processor with multi-format support"""
    async def initialize(self):
        """
        Optional startup logic. Add any preloading or directory scanning here.
        """
        pass
    def __init__(self, settings):
        self.settings = settings
        self.upload_dir = Path(settings.UPLOAD_DIR)
        self.upload_dir.mkdir(exist_ok=True)
        
        # Initialize speech recognition
        self.speech_recognizer = sr.Recognizer()
        
        # Supported file types and processors
        self.processors = {
            # Text files
            '.txt': self._process_text,
            '.md': self._process_text,
            '.rtf': self._process_text,
            
            # Documents
            '.pdf': self._process_pdf,
            '.docx': self._process_docx,
            '.doc': self._process_docx,
            
            # Spreadsheets
            '.xlsx': self._process_excel,
            '.xls': self._process_excel,
            '.csv': self._process_csv,
            
            # Data formats
            '.json': self._process_json,
            '.xml': self._process_xml,
            '.yaml': self._process_yaml,
            '.yml': self._process_yaml,
            
            # Code files
            '.py': self._process_code,
            '.js': self._process_code,
            '.ts': self._process_code,
            '.java': self._process_code,
            '.cpp': self._process_code,
            '.c': self._process_code,
            '.go': self._process_code,
            '.rs': self._process_code,
            '.php': self._process_code,
            '.rb': self._process_code,
            '.html': self._process_html,
            '.css': self._process_code,
            '.sql': self._process_code,
            
            # Images
            '.jpg': self._process_image,
            '.jpeg': self._process_image,
            '.png': self._process_image,
            '.gif': self._process_image,
            '.bmp': self._process_image,
            '.tiff': self._process_image,
            '.svg': self._process_svg,
            
            # Audio
            '.mp3': self._process_audio,
            '.wav': self._process_audio,
            '.m4a': self._process_audio,
            '.flac': self._process_audio,
            '.ogg': self._process_audio,
            
            # Video
            '.mp4': self._process_video,
            '.avi': self._process_video,
            '.mov': self._process_video,
            '.mkv': self._process_video,
            '.wmv': self._process_video,
            
            # Archives
            '.zip': self._process_zip,
            '.tar': self._process_tar,
            '.gz': self._process_tar,
            '.rar': self._process_rar
        }
    
    async def process_file(self, file, project_id: str) -> FileProcessResult:
        """Process uploaded file and extract content"""
        start_time = asyncio.get_event_loop().time()
        
        try:
            # Validate file
            if not self._validate_file(file):
                raise ValueError(f"File type not supported: {file.filename}")
            
            # Save file temporarily
            temp_path = await self._save_temp_file(file, project_id)
            
            # Detect file type
            file_type = self._detect_file_type(temp_path)
            file_ext = Path(file.filename).suffix.lower()
            
            # Process based on file type
            processor = self.processors.get(file_ext, self._process_binary)
            content, metadata = await processor(temp_path, file.filename)
            
            # Calculate processing time
            processing_time = asyncio.get_event_loop().time() - start_time
            
            return FileProcessResult(
                filename=file.filename,
                file_type=file_type,
                size=file.size,
                content=content,
                metadata=metadata,
                processing_time=processing_time,
                success=True
            )
            
        except Exception as e:
            logger.error(f"File processing error for {file.filename}: {str(e)}")
            processing_time = asyncio.get_event_loop().time() - start_time
            
            return FileProcessResult(
                filename=file.filename,
                file_type="unknown",
                size=file.size if hasattr(file, 'size') else 0,
                content="",
                metadata={},
                processing_time=processing_time,
                success=False,
                error=str(e)
            )
        finally:
            # Cleanup temp file
            if 'temp_path' in locals() and temp_path.exists():
                temp_path.unlink()
    
    def _validate_file(self, file) -> bool:
        """Validate uploaded file"""
        if not file.filename:
            return False
        
        # Check file extension
        file_ext = Path(file.filename).suffix.lower()
        if file_ext not in self.settings.ALLOWED_FILE_TYPES:
            return False
        
        # Check file size
        if hasattr(file, 'size') and file.size > self.settings.MAX_FILE_SIZE:
            return False
        
        return True
    
    async def _save_temp_file(self, file, project_id: str) -> Path:
        """Save file to temporary location"""
        temp_dir = Path(self.settings.TEMP_DIR) / project_id
        temp_dir.mkdir(parents=True, exist_ok=True)
        
        temp_path = temp_dir / file.filename
        
        async with aiofiles.open(temp_path, 'wb') as f:
            content = await file.read()
            await f.write(content)
        
        return temp_path
    
    def _detect_file_type(self, file_path: Path) -> str:
        """Detect file MIME type"""
        try:
            mime_type = magic.from_file(str(file_path), mime=True)
            return mime_type
        except:
            return mimetypes.guess_type(str(file_path))[0] or "application/octet-stream"
    
    # Text Processing
    async def _process_text(self, file_path: Path, filename: str) -> tuple[str, Dict[str, Any]]:
        """Process text files"""
        async with aiofiles.open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = await f.read()
        
        metadata = {
            'lines': len(content.split('\n')),
            'words': len(content.split()),
            'characters': len(content),
            'encoding': 'utf-8'
        }
        
        return content, metadata
    
    # PDF Processing
    async def _process_pdf(self, file_path: Path, filename: str) -> tuple[str, Dict[str, Any]]:
        """Process PDF files"""
        content = ""
        metadata = {'pages': 0, 'method': 'text_extraction'}
        
        try:
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                metadata['pages'] = len(pdf_reader.pages)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        content += f"\n--- Page {page_num + 1} ---\n{page_text}\n"
                    except Exception as e:
                        logger.warning(f"Error extracting text from page {page_num + 1}: {str(e)}")
                        continue
        
        except Exception as e:
            logger.error(f"PDF processing error: {str(e)}")
            # Fallback to OCR if text extraction fails
            content, ocr_metadata = await self._process_pdf_ocr(file_path, filename)
            metadata.update(ocr_metadata)
            metadata['method'] = 'ocr_fallback'
        
        return content.strip(), metadata
    
    async def _process_pdf_ocr(self, file_path: Path, filename: str) -> tuple[str, Dict[str, Any]]:
        """Process PDF using OCR"""
        try:
            import fitz  # PyMuPDF
            
            content = ""
            doc = fitz.open(str(file_path))
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                pix = page.get_pixmap()
                img_data = pix.tobytes("png")
                
                # Convert to PIL Image
                from io import BytesIO
                img = Image.open(BytesIO(img_data))
                
                # OCR
                page_text = pytesseract.image_to_string(img)
                content += f"\n--- Page {page_num + 1} ---\n{page_text}\n"
            
            doc.close()
            
            metadata = {
                'pages': len(doc),
                'method': 'ocr',
                'ocr_confidence': 'medium'
            }
            
            return content.strip(), metadata
            
        except Exception as e:
            logger.error(f"PDF OCR processing error: {str(e)}")
            return f"Error processing PDF: {str(e)}", {'error': str(e)}
    
    # Document Processing
    async def _process_docx(self, file_path: Path, filename: str) -> tuple[str, Dict[str, Any]]:
        """Process DOCX files"""
        try:
            doc = docx.Document(str(file_path))
            content = ""
            
            for paragraph in doc.paragraphs:
                content += paragraph.text + "\n"
            
            # Extract tables
            for table in doc.tables:
                content += "\n--- Table ---\n"
                for row in table.rows:
                    row_text = " | ".join([cell.text for cell in row.cells])
                    content += row_text + "\n"
            
            metadata = {
                'paragraphs': len(doc.paragraphs),
                'tables': len(doc.tables),
                'method': 'docx_extraction'
            }
            
            return content.strip(), metadata
            
        except Exception as e:
            logger.error(f"DOCX processing error: {str(e)}")
            return f"Error processing DOCX: {str(e)}", {'error': str(e)}
    
    # Spreadsheet Processing
    async def _process_excel(self, file_path: Path, filename: str) -> tuple[str, Dict[str, Any]]:
        """Process Excel files"""
        try:
            # Read all sheets
            excel_file = pd.ExcelFile(str(file_path))
            content = ""
            
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(str(file_path), sheet_name=sheet_name)
                content += f"\n--- Sheet: {sheet_name} ---\n"
                content += df.to_string(index=False) + "\n"
            
            metadata = {
                'sheets': len(excel_file.sheet_names),
                'sheet_names': excel_file.sheet_names,
                'method': 'pandas_excel'
            }
            
            return content.strip(), metadata
            
        except Exception as e:
            logger.error(f"Excel processing error: {str(e)}")
            return f"Error processing Excel: {str(e)}", {'error': str(e)}
    
    async def _process_csv(self, file_path: Path, filename: str) -> tuple[str, Dict[str, Any]]:
        """Process CSV files"""
        try:
            df = pd.read_csv(str(file_path))
            content = df.to_string(index=False)
            
            metadata = {
                'rows': len(df),
                'columns': len(df.columns),
                'column_names': list(df.columns),
                'method': 'pandas_csv'
            }
            
            return content, metadata
            
        except Exception as e:
            logger.error(f"CSV processing error: {str(e)}")
            return f"Error processing CSV: {str(e)}", {'error': str(e)}
    
    # Data Format Processing
    async def _process_json(self, file_path: Path, filename: str) -> tuple[str, Dict[str, Any]]:
        """Process JSON files"""
        try:
            async with aiofiles.open(file_path, 'r', encoding='utf-8') as f:
                content = await f.read()
            
            # Parse and pretty print
            data = json.loads(content)
            formatted_content = json.dumps(data, indent=2, ensure_ascii=False)
            
            metadata = {
                'size_bytes': len(content),
                'structure_type': type(data).__name__,
                'method': 'json_parse'
            }
            
            return formatted_content, metadata
            
        except Exception as e:
            logger.error(f"JSON processing error: {str(e)}")
            return content if 'content' in locals() else f"Error processing JSON: {str(e)}", {'error': str(e)}
    
    async def _process_xml(self, file_path: Path, filename: str) -> tuple[str, Dict[str, Any]]:
        """Process XML files"""
        try:
            async with aiofiles.open(file_path, 'r', encoding='utf-8') as f:
                content = await f.read()
            
            # Parse XML
            root = ET.fromstring(content)
            
            # Convert to readable format
            formatted_content = self._xml_to_text(root)
            
            metadata = {
                'root_tag': root.tag,
                'elements_count': len(list(root.iter())),
                'method': 'xml_parse'
            }
            
            return formatted_content, metadata
            
        except Exception as e:
            logger.error(f"XML processing error: {str(e)}")
            return content if 'content' in locals() else f"Error processing XML: {str(e)}", {'error': str(e)}
    
    def _xml_to_text(self, element, level=0) -> str:
        """Convert XML element to readable text"""
        indent = "  " * level
        text = f"{indent}<{element.tag}>"
        
        if element.text and element.text.strip():
            text += f" {element.text.strip()}"
        
        text += "\n"
        
        for child in element:
            text += self._xml_to_text(child, level + 1)
        
        return text
    
    async def _process_yaml(self, file_path: Path, filename: str) -> tuple[str, Dict[str, Any]]:
        """Process YAML files"""
        try:
            import yaml
            
            async with aiofiles.open(file_path, 'r', encoding='utf-8') as f:
                content = await f.read()
            
            # Parse YAML
            data = yaml.safe_load(content)
            formatted_content = yaml.dump(data, default_flow_style=False, indent=2)
            
            metadata = {
                'structure_type': type(data).__name__,
                'method': 'yaml_parse'
            }
            
            return formatted_content, metadata
            
        except Exception as e:
            logger.error(f"YAML processing error: {str(e)}")
            return content if 'content' in locals() else f"Error processing YAML: {str(e)}", {'error': str(e)}
    
    # Code Processing
    async def _process_code(self, file_path: Path, filename: str) -> tuple[str, Dict[str, Any]]:
        """Process code files"""
        async with aiofiles.open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = await f.read()
        
        # Detect programming language
        language = self._detect_language(filename)
        
        # Basic code analysis
        lines = content.split('\n')
        non_empty_lines = [line for line in lines if line.strip()]
        comment_lines = [line for line in lines if line.strip().startswith(('#', '//', '/*', '*', '--'))]
        
        metadata = {
            'language': language,
            'total_lines': len(lines),
            'code_lines': len(non_empty_lines),
            'comment_lines': len(comment_lines),
            'method': 'code_analysis'
        }
        
        return content, metadata
    
    def _detect_language(self, filename: str) -> str:
        """Detect programming language from filename"""
        ext_to_lang = {
            '.py': 'python',
            '.js': 'javascript',
            '.ts': 'typescript',
            '.java': 'java',
            '.cpp': 'cpp',
            '.c': 'c',
            '.go': 'go',
            '.rs': 'rust',
            '.php': 'php',
            '.rb': 'ruby',
            '.css': 'css',
            '.sql': 'sql',
            '.html': 'html',
            '.xml': 'xml'
        }
        
        ext = Path(filename).suffix.lower()
        return ext_to_lang.get(ext, 'unknown')
    
    async def _process_html(self, file_path: Path, filename: str) -> tuple[str, Dict[str, Any]]:
        """Process HTML files"""
        async with aiofiles.open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = await f.read()
        
        # Parse HTML and extract text
        soup = BeautifulSoup(content, 'html.parser')
        text_content = soup.get_text(separator='\n', strip=True)
        
        # Extract metadata
        title = soup.find('title')
        meta_tags = soup.find_all('meta')
        
        metadata = {
            'title': title.text if title else 'No title',
            'meta_tags': len(meta_tags),
            'links': len(soup.find_all('a')),
            'images': len(soup.find_all('img')),
            'method': 'html_parse'
        }
        
        return f"HTML Content:\n{text_content}\n\nRaw HTML:\n{content}", metadata
    
    # Image Processing
    async def _process_image(self, file_path: Path, filename: str) -> tuple[str, Dict[str, Any]]:
        """Process image files with OCR"""
        try:
            # Open image
            image = Image.open(str(file_path))
            
            # Basic image info
            width, height = image.size
            mode = image.mode
            format_type = image.format
            
            # OCR text extraction
            ocr_text = pytesseract.image_to_string(image)
            
            # Advanced image analysis
            analysis = await self._analyze_image_content(image)
            
            content = f"Image Analysis:\n"
            content += f"Dimensions: {width}x{height}\n"
            content += f"Mode: {mode}\n"
            content += f"Format: {format_type}\n\n"
            
            if ocr_text.strip():
                content += f"Extracted Text (OCR):\n{ocr_text}\n\n"
            
            content += f"Visual Analysis:\n{analysis}"
            
            metadata = {
                'width': width,
                'height': height,
                'mode': mode,
                'format': format_type,
                'has_text': bool(ocr_text.strip()),
                'method': 'image_ocr_analysis'
            }
            
            return content, metadata
            
        except Exception as e:
            logger.error(f"Image processing error: {str(e)}")
            return f"Error processing image: {str(e)}", {'error': str(e)}
    
    async def _analyze_image_content(self, image: Image.Image) -> str:
        """Analyze image content using computer vision"""
        try:
            # Convert PIL to OpenCV
            cv_image = cv2.cvtColor(np.array(image), cv2.COLOR_RGB2BGR)
            
            # Basic analysis
            analysis = []
            
            # Color analysis
            mean_color = cv2.mean(cv_image)
            analysis.append(f"Average color (BGR): {[int(x) for x in mean_color[:3]]}")
            
            # Brightness analysis
            gray = cv2.cvtColor(cv_image, cv2.COLOR_BGR2GRAY)
            brightness = np.mean(gray)
            analysis.append(f"Average brightness: {brightness:.1f}/255")
            
            # Edge detection
            edges = cv2.Canny(gray, 50, 150)
            edge_density = np.sum(edges > 0) / edges.size
            analysis.append(f"Edge density: {edge_density:.3f}")
            
            # Face detection (basic)
            try:
                face_cascade = cv2.CascadeClassifier(cv2.data.haarcascades + 'haarcascade_frontalface_default.xml')
                faces = face_cascade.detectMultiScale(gray, 1.1, 4)
                analysis.append(f"Detected faces: {len(faces)}")
            except:
                analysis.append("Face detection: unavailable")
            
            return "\n".join(analysis)
            
        except Exception as e:
            return f"Visual analysis error: {str(e)}"
    
    async def _process_svg(self, file_path: Path, filename: str) -> tuple[str, Dict[str, Any]]:
        """Process SVG files"""
        async with aiofiles.open(file_path, 'r', encoding='utf-8') as f:
            content = await f.read()
        
        # Parse SVG
        soup = BeautifulSoup(content, 'xml')
        
        # Extract text elements
        text_elements = soup.find_all('text')
        extracted_text = "\n".join([elem.get_text() for elem in text_elements])
        
        metadata = {
            'text_elements': len(text_elements),
            'total_elements': len(soup.find_all()),
            'method': 'svg_parse'
        }
        
        result = f"SVG Content Analysis:\n"
        if extracted_text:
            result += f"Extracted Text:\n{extracted_text}\n\n"
        result += f"Raw SVG:\n{content}"
        
        return result, metadata
    
    # Audio Processing
    async def _process_audio(self, file_path: Path, filename: str) -> tuple[str, Dict[str, Any]]:
        """Process audio files with speech recognition"""
        try:
            # Load audio
            audio = AudioSegment.from_file(str(file_path))
            
            # Convert to WAV for speech recognition
            wav_path = file_path.with_suffix('.wav')
            audio.export(str(wav_path), format="wav")
            
            # Speech recognition
            transcript = ""
            try:
                with sr.AudioFile(str(wav_path)) as source:
                    audio_data = self.speech_recognizer.record(source)
                    transcript = self.speech_recognizer.recognize_google(audio_data)
            except Exception as e:
                transcript = f"Speech recognition failed: {str(e)}"
            
            # Audio analysis
            duration = len(audio) / 1000.0  # seconds
            sample_rate = audio.frame_rate
            channels = audio.channels
            
            content = f"Audio Analysis:\n"
            content += f"Duration: {duration:.2f} seconds\n"
            content += f"Sample Rate: {sample_rate} Hz\n"
            content += f"Channels: {channels}\n\n"
            content += f"Transcript:\n{transcript}"
            
            metadata = {
                'duration': duration,
                'sample_rate': sample_rate,
                'channels': channels,
                'has_transcript': bool(transcript and not transcript.startswith("Speech recognition failed")),
                'method': 'audio_speech_recognition'
            }
            
            # Cleanup temp WAV file
            if wav_path.exists():
                wav_path.unlink()
            
            return content, metadata
            
        except Exception as e:
            logger.error(f"Audio processing error: {str(e)}")
            return f"Error processing audio: {str(e)}", {'error': str(e)}
    
    # Video Processing
    async def _process_video(self, file_path: Path, filename: str) -> tuple[str, Dict[str, Any]]:
        """Process video files"""
        try:
            # Use OpenCV to analyze video
            cap = cv2.VideoCapture(str(file_path))
            
            # Get video properties
            fps = cap.get(cv2.CAP_PROP_FPS)
            frame_count = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))
            duration = frame_count / fps if fps > 0 else 0
            width = int(cap.get(cv2.CAP_PROP_FRAME_WIDTH))
            height = int(cap.get(cv2.CAP_PROP_FRAME_HEIGHT))
            
            # Extract frames for analysis
            frames_analyzed = 0
            frame_descriptions = []
            
            # Analyze every 30th frame (or fewer if short video)
            frame_interval = max(1, frame_count // 10)
            
            for i in range(0, frame_count, frame_interval):
                cap.set(cv2.CAP_PROP_POS_FRAMES, i)
                ret, frame = cap.read()
                
                if ret and frames_analyzed < 5:  # Limit analysis
                    # Basic frame analysis
                    timestamp = i / fps
                    brightness = np.mean(cv2.cvtColor(frame, cv2.COLOR_BGR2GRAY))
                    frame_descriptions.append(f"Frame at {timestamp:.1f}s: brightness={brightness:.1f}")
                    frames_analyzed += 1
            
            cap.release()
            
            content = f"Video Analysis:\n"
            content += f"Duration: {duration:.2f} seconds\n"
            content += f"Resolution: {width}x{height}\n"
            content += f"FPS: {fps:.2f}\n"
            content += f"Total Frames: {frame_count}\n\n"
            content += "Frame Analysis:\n" + "\n".join(frame_descriptions)
            
            metadata = {
                'duration': duration,
                'width': width,
                'height': height,
                'fps': fps,
                'frame_count': frame_count,
                'frames_analyzed': frames_analyzed,
                'method': 'video_opencv_analysis'
            }
            
            return content, metadata
            
        except Exception as e:
            logger.error(f"Video processing error: {str(e)}")
            return f"Error processing video: {str(e)}", {'error': str(e)}
    
    # Archive Processing
    async def _process_zip(self, file_path: Path, filename: str) -> tuple[str, Dict[str, Any]]:
        """Process ZIP archives"""
        try:
            content = "ZIP Archive Contents:\n\n"
            file_list = []
            
            with zipfile.ZipFile(str(file_path), 'r') as zip_file:
                for file_info in zip_file.filelist:
                    file_list.append({
                        'name': file_info.filename,
                        'size': file_info.file_size,
                        'compressed_size': file_info.compress_size,
                        'date': f"{file_info.date_time[0]}-{file_info.date_time[1]:02d}-{file_info.date_time[2]:02d}"
                    })
                    
                    content += f"File: {file_info.filename}\n"
                    content += f"  Size: {file_info.file_size} bytes\n"
                    content += f"  Compressed: {file_info.compress_size} bytes\n"
                    content += f"  Date: {file_info.date_time[0]}-{file_info.date_time[1]:02d}-{file_info.date_time[2]:02d}\n\n"
            
            metadata = {
                'total_files': len(file_list),
                'file_list': file_list,
                'method': 'zip_analysis'
            }
            
            return content, metadata
            
        except Exception as e:
            logger.error(f"ZIP processing error: {str(e)}")
            return f"Error processing ZIP: {str(e)}", {'error': str(e)}
    
    async def _process_tar(self, file_path: Path, filename: str) -> tuple[str, Dict[str, Any]]:
        """Process TAR archives"""
        try:
            content = "TAR Archive Contents:\n\n"
            file_list = []
            
            with tarfile.open(str(file_path), 'r') as tar_file:
                for member in tar_file.getmembers():
                    file_list.append({
                        'name': member.name,
                        'size': member.size,
                        'type': 'directory' if member.isdir() else 'file'
                    })
                    
                    content += f"{'Directory' if member.isdir() else 'File'}: {member.name}\n"
                    content += f"  Size: {member.size} bytes\n\n"
            
            metadata = {
                'total_items': len(file_list),
                'file_list': file_list,
                'method': 'tar_analysis'
            }
            
            return content, metadata
            
        except Exception as e:
            logger.error(f"TAR processing error: {str(e)}")
            return f"Error processing TAR: {str(e)}", {'error': str(e)}
    
    async def _process_rar(self, file_path: Path, filename: str) -> tuple[str, Dict[str, Any]]:
        """Process RAR archives"""
        try:
            import rarfile
            
            content = "RAR Archive Contents:\n\n"
            file_list = []
            
            with rarfile.RarFile(str(file_path)) as rar_file:
                for file_info in rar_file.infolist():
                    file_list.append({
                        'name': file_info.filename,
                        'size': file_info.file_size,
                        'compressed_size': file_info.compress_size
                    })
                    
                    content += f"File: {file_info.filename}\n"
                    content += f"  Size: {file_info.file_size} bytes\n"
                    content += f"  Compressed: {file_info.compress_size} bytes\n\n"
            
            metadata = {
                'total_files': len(file_list),
                'file_list': file_list,
                'method': 'rar_analysis'
            }
            
            return content, metadata
            
        except Exception as e:
            logger.error(f"RAR processing error: {str(e)}")
            return f"Error processing RAR: {str(e)}", {'error': str(e)}
    
    # Binary Processing
    async def _process_binary(self, file_path: Path, filename: str) -> tuple[str, Dict[str, Any]]:
        """Process binary files"""
        try:
            file_size = file_path.stat().st_size
            
            # Read first 1KB for analysis
            with open(file_path, 'rb') as f:
                header = f.read(1024)
            
            # Try to detect file type from header
            file_type = magic.from_buffer(header, mime=True)
            
            content = f"Binary File Analysis:\n"
            content += f"File Size: {file_size} bytes\n"
            content += f"Detected Type: {file_type}\n"
            content += f"Header (hex): {header[:100].hex()}\n"
            
            # Try to find readable strings
            readable_strings = []
            try:
                text_content = header.decode('utf-8', errors='ignore')
                words = [word for word in text_content.split() if len(word) > 3 and word.isalnum()]
                readable_strings = words[:10]  # First 10 readable words
            except:
                pass
            
            if readable_strings:
                content += f"Readable strings: {', '.join(readable_strings)}\n"
            
            metadata = {
                'file_size': file_size,
                'detected_type': file_type,
                'has_readable_content': bool(readable_strings),
                'method': 'binary_analysis'
            }
            
            return content, metadata
            
        except Exception as e:
            logger.error(f"Binary processing error: {str(e)}")
            return f"Error processing binary file: {str(e)}", {'error': str(e)}
    
    async def health_check(self) -> Dict[str, Any]:
        """Health check for file processor"""
        return {
            'status': 'healthy',
            'supported_formats': len(self.processors),
            'upload_dir': str(self.upload_dir),
            'processors_available': list(self.processors.keys())
        }
